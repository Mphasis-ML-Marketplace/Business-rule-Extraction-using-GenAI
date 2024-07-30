import os
import boto3
import json
import zipfile
from docx import Document
import PyPDF2
from botocore.exceptions import ClientError


def call_claude_sonnet_new(messages, system):
    try:
        with open(os.path.join(os.getcwd(), "Data/Input_data/credentials.json"), 'r') as f:
            credentials = json.load(f)

        # Extract the keys
        aws_access_key_id = credentials['aws_access_key_id']
        aws_secret_access_key = credentials['aws_secret_access_key']
        region_name = credentials['region_name']

        brt = boto3.client(
              service_name='bedrock-runtime',
              region_name=region_name,
              aws_access_key_id=aws_access_key_id,
              aws_secret_access_key=aws_secret_access_key  
        )
        model_id = "anthropic.claude-3-5-sonnet-20240620-v1:0"
        native_request = {
              "anthropic_version": "bedrock-2023-05-31",
              "max_tokens": 2048,
              "temperature": 0.1,
              "system": system,
              "messages": messages,
        }
        request = json.dumps(native_request)
        try:
           response = brt.invoke_model(modelId=model_id, body=request)

        except (ClientError, Exception) as e:
           print(f"ERROR: Can't invoke '{model_id}'. Reason: {e}")
           exit(1)

        model_response = json.loads(response["body"].read())
        response_text = model_response["content"][0]["text"]
        return response_text
    
    except Exception as e:
        return '111'

def rules_advocate(document_text):

    system= "You are a policy claims validation officer. Given the content of a policy document, your job is to identify the major conditions mentioned in the document to process claims. So when a new claim is made, the details of the claim has to be validated against the important rules you have extracted from the policy document. The content of the document is mentioned below.\n\n<Document content>\n" + document_text + "\n<End of document content>"
                
    messages= [
        {"role": "user",
         "content": [
             {
                  "type": "text",
                  "text": "Extract the important conditions mentioned in the document that will help validate a claim made by someone."
             }
          ]
        }
             
    ]
    
    response_text = call_claude_sonnet_new(messages, system)

    return response_text
        
 
 
def supervisor(document_text, document_rules):

    system= "You are a document rules verification officer. You are given the contents of a document and the set of rules extracted from the document for user queries and claims. Your job is to check whether the extracted conditions from the document pertaining to the user claims are all covered or if any conditions are missing from the set of rules.\n\n<Document content>\n" + document_text + "\n<End of document content>\n\nBelow, the extracted set of rules from the document pertaining to user claims are given. \n<Set of extracted rules>\n" + document_rules + "\n<End of rules>\n\nGiven the policy document and the set of rules, If there are any points missing from the set of rules, fix it and output the reformed set of rules pertaining to the user claims. If you think there are any redundant rules, remove them and output the new set of rules. If there is no changes to be made, then output the rule set as it is."

    messages= [
        {"role": "user",
         "content": [
             {
                  "type": "text",
                  "text": "Verify the given set of rules against the input document and fix them if necessary."
             }
          ]
        }
             
    ]
    
    response_text = call_claude_sonnet_new(messages, system)
    return response_text

def extract_text_from_docx(config_file_path):
    document = Document(config_file_path)
    full_text = []
    for paragraph in document.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

def extract_text_from_pdf(config_file_path):
    with open(config_file_path, 'rb') as file:
        reader = PyPDF2.PdfFileReader(file)
        full_text = []
        for page_num in range(reader.numPages):
            page = reader.getPage(page_num)
            full_text.append(page.extractText())
    return '\n'.join(full_text)

def save_to_docx(document_rules, supervisor_output, file_name):     
    doc = Document()   
    doc.add_heading('Generated Rules', level=1)    
    doc.add_paragraph(document_rules)  

    doc.add_heading('Optimized Rules', level=1)
    doc.add_paragraph(supervisor_output) 
        
    doc.save(file_name)
    print(f'DOCX file saved as {file_name}')
    return file_name

def master(zip_path):
    file_names_list = []
    pwd = os.getcwd()
    data_handle = os.path.join(pwd,'Data')
    with open(zip_path, 'rb') as fh:
        z = zipfile.ZipFile(fh)
        for name in z.namelist():
           z.extract(name, data_handle)
           file_names_list.append(os.path.join(data_handle, name))
        fh.close()
    return file_names_list

def main_inference(zip_path):
 
    file_names = master(zip_path)
    output = []
 
    try:
        for name in file_names:
            
            if name.endswith(".docx"):
                config_file_path=name
                document_text = extract_text_from_docx(config_file_path)
            elif name.endswith(".pdf"):
                config_file_path=name
                document_text = extract_text_from_pdf(config_file_path)
            elif os.path.isfile(name):
                continue
            elif os.path.isdir(name):
                document_folder_path=name
                files = os.listdir(document_folder_path)
                if len(files) < 9:
                    for file in files:
                         file_path = os.path.join(document_folder_path, file)
                    if file.endswith(".docx"):
                         document_text = extract_text_from_docx(file_path)
                    elif file.endswith(".pdf"):
                         document_text = extract_text_from_pdf(file_path)
                    else:
                         continue

                    try:
                        document_rules = rules_advocate(document_text)
                        supervisor_output = supervisor(document_rules, document_text)
                        file_name = os.path.basename(name)
                        output_file = save_to_docx(document_rules, supervisor_output, file_name)
                        output.append(output_file)

                    except Exception as e:
                        print(f"An error occurred while processing {file_path}: {e}")
                continue  # Skip the remaining part of the loop for directories

            if 'document_text' in locals():
               try:
                     document_rules = rules_advocate(document_text)
                     supervisor_output = supervisor(document_rules, document_text)
                     file_name = os.path.basename(name)
                     output_file = save_to_docx(document_rules, supervisor_output, file_name)
                     output.append(output_file)

               except Exception as e:
                     print(f"An error occurred: {e}")
    
    except Exception as general_exception:
        print(f"A general error occurred during processing: {general_exception}")

    return output

        
        